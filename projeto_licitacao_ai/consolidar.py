from docx import Document
from docx.shared import Pt
import os


def gerar_documento_final():
    doc = Document()

    # Cabeçalho "Oficial"
    titulo = doc.add_heading('RELATÓRIO CONSOLIDADO DE INSTRUÇÃO PROCESSUAL', 0)
    doc.add_paragraph('Projeto: Auditoria e Elaboração de ETP - Limpeza e Conservação')
    doc.add_paragraph('Gerado automaticamente via IA Agêntica Local')
    doc.add_paragraph('-' * 50)

    # Lista de arquivos para consolidar
    relatorios = [
        ("1. ANÁLISE COMPARATIVA DE MERCADO", "relatorio_analista.md"),
        ("2. MINUTA DE ESTUDO TÉCNICO PRELIMINAR (ETP)", "relatorio_consultor.md"),
        ("3. RELATÓRIO DE AUDITORIA E CONFORMIDADE", "relatorio_auditoria_teste.md")
    ]

    for secao_titulo, nome_arquivo in relatorios:
        doc.add_heading(secao_titulo, level=1)

        if os.path.exists(nome_arquivo):
            with open(nome_arquivo, 'r', encoding='utf-8') as f:
                conteudo = f.read()
                # Removemos as marcações de Markdown simples para o Word não ficar estranho
                limpo = conteudo.replace('#', '').replace('**', '')
                doc.add_paragraph(limpo)
        else:
            doc.add_paragraph(f"Aviso: O arquivo {nome_arquivo} não foi encontrado.")

    # Salvando o resultado
    nome_final = "PARECER_TECNICO_FINAL.docx"
    doc.save(nome_final)
    print(f"\n✨ Sucesso! Documento '{nome_final}' gerado com sucesso.")


# No final do consolidar.py, use sempre isso:
if __name__ == "__main__":
    gerar_documento_final()